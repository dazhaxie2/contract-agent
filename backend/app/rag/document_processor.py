"""
文档预处理管线 - 解析、清洗、元数据提取
支持PDF/DOCX/TXT等多格式，保留法律文档层级结构
"""

import hashlib
import re
import uuid
from pathlib import Path
from typing import BinaryIO

from loguru import logger

from app.core.config import settings


class DocumentProcessor:
    """文档预处理器"""

    SUPPORTED_FORMATS = {".pdf", ".docx", ".doc", ".txt", ".md", ".html"}

    async def process(self, file: BinaryIO, filename: str, tenant_id: str) -> dict:
        """
        完整文档处理流程:
        1. 格式解析 -> 2. 文本清洗 -> 3. 结构识别 -> 4. 元数据提取
        """
        ext = Path(filename).suffix.lower()
        if ext not in self.SUPPORTED_FORMATS:
            raise ValueError(f"不支持的文件格式: {ext}")

        content = file.read()
        file_hash = hashlib.md5(content).hexdigest()

        # Step 1: 格式解析
        raw_text, structure = await self._parse_document(content, ext)

        # Step 2: 文本清洗
        cleaned_text = self._clean_text(raw_text)

        # Step 3: 结构识别
        sections = self._identify_structure(cleaned_text, structure)

        # Step 4: 元数据提取
        metadata = await self._extract_metadata(cleaned_text, filename)

        return {
            "file_hash": file_hash,
            "file_size": len(content),
            "raw_text": raw_text,
            "cleaned_text": cleaned_text,
            "sections": sections,
            "metadata": metadata,
        }

    async def _parse_document(self, content: bytes, ext: str) -> tuple[str, list]:
        """多格式文档解析"""
        if ext == ".pdf":
            return self._parse_pdf(content)
        elif ext in (".docx", ".doc"):
            return self._parse_docx(content)
        elif ext == ".txt":
            return content.decode("utf-8", errors="ignore"), []
        elif ext == ".md":
            return content.decode("utf-8", errors="ignore"), []
        return content.decode("utf-8", errors="ignore"), []

    def _parse_pdf(self, content: bytes) -> tuple[str, list]:
        try:
            import fitz
            doc = fitz.open(stream=content, filetype="pdf")
            full_text = []
            structure = []

            for page_num, page in enumerate(doc):
                text = page.get_text("text")
                full_text.append(text)

                blocks = page.get_text("dict")["blocks"]
                for block in blocks:
                    if "lines" not in block:
                        continue
                    for line in block["lines"]:
                        for span in line["spans"]:
                            if span["size"] > 14:
                                structure.append({
                                    "type": "heading",
                                    "text": span["text"].strip(),
                                    "page": page_num + 1,
                                    "font_size": span["size"],
                                })
            doc.close()

            extracted_text = "\n".join(full_text)
            if len(extracted_text.strip()) < settings.multimodal.min_text_length_for_ocr:
                logger.info("Extracted text too short, attempting OCR fallback")
                doc = fitz.open(stream=content, filetype="pdf")
                ocr_text = self._ocr_pdf_pages(doc)
                doc.close()
                if ocr_text and len(ocr_text.strip()) > len(extracted_text.strip()):
                    extracted_text = ocr_text

            return extracted_text, structure
        except ImportError:
            logger.warning("PyMuPDF未安装，使用纯文本提取")
            return content.decode("utf-8", errors="ignore"), []

    def _ocr_pdf_pages(self, doc) -> str:
        try:
            import numpy as np
            from paddleocr import PaddleOCR
        except ImportError:
            logger.warning("PaddleOCR or numpy not available for OCR fallback")
            return ""

        try:
            ocr = PaddleOCR(use_angle_cls=True, lang="ch", show_log=False)
            all_text: list[str] = []

            for page in doc:
                pix = page.get_pixmap(dpi=200)
                img_bytes = pix.tobytes("png")
                nparr = np.frombuffer(img_bytes, np.uint8)

                try:
                    import cv2
                    img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                except ImportError:
                    img = nparr

                result = ocr.ocr(img, cls=True)
                if result and result[0]:
                    for line in result[0]:
                        if line and len(line) >= 2:
                            text = line[1][0] if isinstance(line[1], (list, tuple)) else str(line[1])
                            all_text.append(text)

            return "\n".join(all_text)
        except Exception as exc:
            logger.error(f"OCR fallback failed: {exc}")
            return ""

    def _parse_docx(self, content: bytes) -> tuple[str, list]:
        """DOCX解析 - python-docx"""
        try:
            import io
            from docx import Document
            doc = Document(io.BytesIO(content))
            full_text = []
            structure = []

            for para in doc.paragraphs:
                full_text.append(para.text)
                if para.style.name.startswith("Heading"):
                    level = int(para.style.name.replace("Heading ", "") or "1")
                    structure.append({
                        "type": f"h{level}",
                        "text": para.text.strip(),
                        "level": level,
                    })

            # 提取表格
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                full_text.append("\n".join(table_text))

            return "\n".join(full_text), structure
        except ImportError:
            return content.decode("utf-8", errors="ignore"), []

    def _clean_text(self, text: str) -> str:
        """文本清洗 - 去噪、归一化"""
        # 去除多余空白
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]{2,}', ' ', text)

        # 去除页眉页脚常见模式
        text = re.sub(r'第\s*\d+\s*页\s*共\s*\d+\s*页', '', text)
        text = re.sub(r'- \d+ -', '', text)

        # 去除水印
        text = re.sub(r'(?i)(confidential|draft|watermark|仅供参考|内部资料)', '', text)

        # Unicode归一化
        text = text.replace('\u3000', ' ')  # 全角空格
        text = text.replace('\xa0', ' ')  # 不间断空格

        return text.strip()

    def _identify_structure(self, text: str, structure_hints: list) -> list[dict]:
        """识别法律文档层级结构: 编/章/节/条/款/项"""
        sections = []

        # 法律条文结构模式
        patterns = [
            (r'^(第[一二三四五六七八九十百千]+编)\s*(.+)$', 'part', 1),
            (r'^(第[一二三四五六七八九十百千]+章)\s*(.+)$', 'chapter', 2),
            (r'^(第[一二三四五六七八九十百千]+节)\s*(.+)$', 'section', 3),
            (r'^(第[一二三四五六七八九十百千万]+条)\s*(.*)$', 'article', 4),
            (r'^(（[一二三四五六七八九十]+）)\s*(.*)$', 'paragraph', 5),
            (r'^([一二三四五六七八九十]+、)\s*(.*)$', 'clause', 5),
            (r'^(\d+[.、])\s*(.+)$', 'item', 6),
        ]

        for line in text.split('\n'):
            line = line.strip()
            if not line:
                continue

            for pattern, level_name, level_num in patterns:
                match = re.match(pattern, line)
                if match:
                    sections.append({
                        "id": str(uuid.uuid4()),
                        "type": level_name,
                        "level": level_num,
                        "number": match.group(1),
                        "title": match.group(2).strip() if match.lastindex >= 2 else "",
                        "content": line,
                    })
                    break

        return sections

    async def _extract_metadata(self, text: str, filename: str) -> dict:
        """元数据提取 - 文档类型/发布机构/生效日期等"""
        metadata = {
            "doc_type": self._classify_doc_type(text, filename),
            "title": self._extract_title(text, filename),
        }

        # 日期提取
        date_patterns = [
            (r'(\d{4})年(\d{1,2})月(\d{1,2})日.*(?:起)?施行', 'effective_date'),
            (r'(\d{4})年(\d{1,2})月(\d{1,2})日.*(?:起)?生效', 'effective_date'),
            (r'(\d{4})年(\d{1,2})月(\d{1,2})日.*发布', 'publish_date'),
            (r'自(\d{4})年(\d{1,2})月(\d{1,2})日起.*废止', 'expiry_date'),
        ]
        for pattern, field in date_patterns:
            match = re.search(pattern, text)
            if match:
                metadata[field] = f"{match.group(1)}-{match.group(2).zfill(2)}-{match.group(3).zfill(2)}"

        # 发布机构提取
        authority_patterns = [
            r'(全国人民代表大会[常务委员会]*)',
            r'(国务院)',
            r'(最高人民法院)',
            r'(最高人民检察院)',
            r'([^\s]{2,20}(?:部|委|局|厅|办|院|会))',
        ]
        for pattern in authority_patterns:
            match = re.search(pattern, text[:500])
            if match:
                metadata["issuing_authority"] = match.group(1)
                break

        return metadata

    @staticmethod
    def _classify_doc_type(text: str, filename: str) -> str:
        """文档类型分类"""
        text_lower = text[:2000].lower()
        fn_lower = filename.lower()

        if any(kw in text_lower for kw in ["中华人民共和国", "法律", "条例", "办法", "规定"]):
            return "law"
        if any(kw in fn_lower for kw in ["合同", "协议", "contract"]):
            return "contract"
        if any(kw in text_lower for kw in ["合规", "规范", "标准", "指引"]):
            return "regulation"
        if any(kw in text_lower for kw in ["判决书", "裁定书", "案例"]):
            return "case"
        return "guide"

    @staticmethod
    def _extract_title(text: str, filename: str) -> str:
        """提取文档标题"""
        lines = text.strip().split('\n')
        for line in lines[:10]:
            line = line.strip()
            if len(line) > 5 and len(line) < 100:
                return line
        return Path(filename).stem


document_processor = DocumentProcessor()
