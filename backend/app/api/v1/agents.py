"""Agent execution APIs backed by relational persistence."""

from __future__ import annotations

import json
import io
import time
import uuid
from datetime import datetime, timedelta, timezone
from typing import Any

from fastapi import APIRouter, Depends, HTTPException, Query, Request
from fastapi.responses import StreamingResponse
from loguru import logger
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.agents.orchestrator import orchestrator_agent
from app.agents.tool_catalog import tool_catalog_by_domain
from app.core.database import get_read_db, get_write_db
from app.core.doc_types import ENTERPRISE_RULE, basis_kind
from app.core.request_context import RequestContext, resolve_request_context
from app.models.agent import AgentDecision, AgentExecution, AgentStep
from app.models.retrieval import RetrievalLog
from app.rag.context_builder import context_builder
from app.rag.retriever import hybrid_retriever
from app.schemas.agent import (
    AgentDecisionExecuteRequest,
    AgentExecuteRequest,
    AgentExecuteResponse,
    AgentPlanRequest,
    AgentPlanResponse,
    PlanStep,
)
from app.schemas.contract_workspace import ContractMatter, ExpenseEvidence, PerformanceEvent
from app.services.citation_service import citation_service
from app.services.llm_service import llm_service
from app.services.session_memory_service import session_memory_service

router = APIRouter()

DECISION_TTL_SECONDS = 60 * 60
ENTERPRISE_RULE_TOP_K = 4


def get_request_context(request: Request) -> RequestContext:
    return resolve_request_context(request)


def _parse_uuid(value: str, field_name: str) -> uuid.UUID:
    try:
        return uuid.UUID(value)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=f"invalid {field_name}") from exc


def _step_payload(step: AgentStep) -> dict:
    content = step.observation or step.thought or step.action or ""
    return {
        "step_id": str(step.id),
        "span_id": step.span_id,
        "parent_span_id": step.parent_span_id,
        "step_number": step.step_number,
        "step_type": step.step_type,
        "content": content,
        "action": step.action,
        "tool_name": step.tool_name,
        "status": step.status,
        "tokens_used": step.tokens_used,
        "latency_ms": round(float(step.latency_ms or 0.0), 2),
    }


def _execution_payload(execution: AgentExecution, steps: list[dict] | None = None) -> dict:
    metadata = execution.result_metadata or {}
    usage = metadata.get("usage") or {}
    if "total_tokens" not in usage:
        usage["total_tokens"] = execution.total_tokens_used
    return {
        "execution_id": execution.id,
        "trace_id": execution.trace_id,
        "status": execution.status,
        "result": execution.result or "",
        "references": metadata.get("references", []),
        "review_report": metadata.get("review_report"),
        "decision_id": metadata.get("decision_id"),
        "plan": metadata.get("plan"),
        "tool_results": metadata.get("tool_results", []),
        "regression_case_id": metadata.get("regression_case_id"),
        "user_feedback": execution.user_feedback,
        "steps": steps or [],
        "usage": usage,
        "latency_ms": round(float(execution.latency_ms or 0.0), 2),
        "task_type": execution.task_type,
        "created_at": execution.created_at,
        "completed_at": execution.completed_at,
    }


def _severity_from_text(text: str) -> str:
    lowered = text.lower()
    if any(marker in text for marker in ["高风险", "严重", "重大"]) or any(
        marker in lowered for marker in ["high risk", "critical", "severe"]
    ):
        return "high"
    if any(marker in text for marker in ["低风险", "轻微"]) or any(marker in lowered for marker in ["low risk", "minor"]):
        return "low"
    return "medium"


def _first_sentence(text: str, limit: int = 360) -> str:
    normalized = " ".join((text or "").strip().split())
    if len(normalized) <= limit:
        return normalized
    return normalized[:limit].rstrip() + "..."


def _basis_join(refs: list[dict]) -> str:
    return "；".join(
        str(item.get("citation_code") or item.get("doc_title") or item.get("chunk_id"))
        for item in refs
        if item.get("citation_code") or item.get("doc_title") or item.get("chunk_id")
    )


def _build_contract_review_report(result_text: str, references: list[dict], query: str) -> dict[str, Any]:
    citation_refs = [
        {
            "ref_id": ref.get("ref_id"),
            "citation_id": ref.get("citation_id"),
            "citation_code": ref.get("citation_code"),
            "doc_title": ref.get("doc_title"),
            "hierarchy": ref.get("hierarchy"),
            "chunk_id": ref.get("chunk_id"),
            "doc_type": ref.get("doc_type"),
            "basis_kind": basis_kind(ref.get("doc_type")),
        }
        for ref in references
        if ref.get("citation_id") or ref.get("citation_code")
    ]
    has_citation = bool(citation_refs)
    legal_refs = [item for item in citation_refs if item["basis_kind"] == "legal"]
    enterprise_refs = [item for item in citation_refs if item["basis_kind"] == "enterprise"]
    severity = _severity_from_text(result_text)
    confidence = 0.72 if has_citation else 0.35
    risk_title = "合同审查结论" if has_citation else "不确定：缺少可追溯依据"

    # 法律依据优先取权威法律来源；若没有命中法律来源，回退到全部引用避免空依据
    legal_basis = _basis_join((legal_refs or citation_refs)[:5]) or "未检索到可验证引用依据"
    enterprise_basis = _basis_join(enterprise_refs[:5])

    return {
        "overall_risk": severity if has_citation else "uncertain",
        "summary": _first_sentence(result_text, limit=500) or "未生成有效审查结论",
        "risk_items": [
            {
                "severity": severity if has_citation else "uncertain",
                "clause_excerpt": _first_sentence(query, limit=320),
                "issue": risk_title,
                "legal_basis": legal_basis,
                "enterprise_basis": enterprise_basis,
                "recommendation": _first_sentence(result_text, limit=600)
                or "建议补充法规依据后再形成正式审查意见",
                "confidence": confidence,
                "references": citation_refs[:8],
            }
        ],
        "enterprise_rule_count": len(enterprise_refs),
        "generated_from": "agent_execution",
    }


async def _expire_stale_decisions(db: AsyncSession, now: datetime | None = None) -> None:
    current = now or datetime.now(timezone.utc)
    rows = (
        await db.scalars(
            select(AgentDecision).where(
                AgentDecision.status == "planned",
                AgentDecision.expires_at < current,
            )
        )
    ).all()
    for row in rows:
        row.status = "expired"
        row.updated_at = current


def _collect_document_filters(context: dict, filters: dict) -> dict:
    merged = dict(filters or {})
    doc_id = context.get("doc_id") or context.get("document_id") or merged.get("doc_id")
    document_ids = context.get("document_ids") or merged.get("document_ids")
    if doc_id:
        merged["doc_id"] = str(doc_id)
    if document_ids:
        merged["document_ids"] = [str(item) for item in document_ids]
    elif doc_id:
        merged["document_ids"] = [str(doc_id)]
    return merged


async def _merge_enterprise_rules(*, query: str, tenant_id: str, base_results: list) -> list:
    """额外检索企业自有制度并并入审查上下文。

    企业制度面向整个租户知识库，不受当前合同 doc_id 过滤约束，
    因此独立发起一次按 doc_type 限定的检索，确保制度依据有机会进入上下文，
    而不是和法规在同一次召回里相互挤占。
    """
    try:
        rule_hits = await hybrid_retriever.retrieve(
            query=query,
            tenant_id=tenant_id,
            filters={"doc_type": ENTERPRISE_RULE, "effective": True},
            top_k=ENTERPRISE_RULE_TOP_K,
        )
    except Exception as exc:
        logger.debug(f"enterprise rule retrieval skipped tenant={tenant_id}: {exc}")
        return base_results
    if not rule_hits:
        return base_results

    seen = {item.chunk_id for item in base_results}
    merged = list(base_results)
    for hit in rule_hits:
        if hit.chunk_id in seen:
            continue
        seen.add(hit.chunk_id)
        merged.append(hit)
    return merged


def _step(
    step_id: str,
    title: str,
    description: str,
    domain: str,
    tool: str,
    *,
    action: str = "read",
    mutates_state: bool = False,
) -> PlanStep:
    return PlanStep(
        step_id=step_id,
        title=title,
        description=description,
        domain=domain,
        tool=tool,
        action=action,
        mutates_state=mutates_state,
    )


def _build_agent_plan(req: AgentPlanRequest) -> dict[str, Any]:
    query = req.query.strip()
    context = dict(req.context or {})
    filters = _collect_document_filters(context, req.filters)
    doc_title = str(context.get("doc_title") or context.get("title") or "当前合同")
    review_type = str(context.get("review_type") or "general")

    steps: list[PlanStep] = []
    if not (filters.get("doc_id") or filters.get("document_ids")):
        steps.append(
            _step(
                "select_contract",
                "确认合同范围",
                "用户尚未提供明确合同，执行前需要选择或上传待审合同。",
                "contract",
                "document.select",
            )
        )
    else:
        steps.append(
            _step(
                "read_contract",
                "读取合同与片段",
                f"读取《{doc_title}》的文档元数据、入库状态和完整 chunk。",
                "contract",
                "document.read",
            )
        )

    steps.extend(
        [
            _step(
                "extract_clauses",
                "提取重点条款",
                "围绕付款、违约、解除、保密、责任限制、争议解决等条款建立审查上下文。",
                "contract",
                "clause.extract",
            ),
            _step(
                "retrieve_legal_basis",
                "检索法规与制度依据",
                "在知识域检索法规、案例、企业制度和既有合同知识，并保留可点击引用。",
                "knowledge",
                "retrieval.search",
            ),
            _step(
                "review_risks",
                "识别合规风险",
                f"按 {review_type} 场景识别条款合法性、履约风险和低置信度问题。",
                "review",
                "compliance.review",
            ),
            _step(
                "draft_recommendations",
                "生成修改建议",
                "输出风险说明、法律依据、可执行修改建议和 Markdown 审查报告。",
                "review",
                "drafting.suggest",
            ),
        ]
    )

    wants_export = any(keyword in query for keyword in ["导出", "下载", "markdown", "Markdown", "报告"])
    if wants_export:
        steps.append(
            _step(
                "export_markdown",
                "准备 Markdown 导出",
                "把审查结论整理成可复制或下载的 Markdown 内容。",
                "review",
                "report.export_markdown",
            )
        )

    if any(keyword in query for keyword in ["反馈", "有问题", "修正", "回归"]):
        steps.append(
            _step(
                "capture_regression_case",
                "预留问题反馈样例",
                "若用户标记结果有问题，保存输入、计划、工具轨迹、引用上下文和期望修正。",
                "observability",
                "regression.capture",
                action="write",
                mutates_state=True,
            )
        )

    estimated_changes = [
        "写入一次 AgentExecution 审查记录",
        "保存 decision_id、执行计划、工具轨迹和 review_report 到 result_metadata",
        "追加会话消息，便于后续追问和历史回放",
    ]
    requires_confirmation = True
    intent_summary = _first_sentence(query, limit=160) or "合同审查请求"

    return {
        "intent_summary": intent_summary,
        "steps": [step.model_dump() for step in steps],
        "requires_confirmation": requires_confirmation,
        "estimated_changes": estimated_changes,
        "context": {
            **context,
            "task_type": req.task_type,
            "filters": filters,
            "review_type": review_type,
        },
    }


def _decision_to_plan_response(record: AgentDecision) -> AgentPlanResponse:
    return AgentPlanResponse(
        decision_id=record.decision_id,
        intent_summary=record.intent_summary,
        steps=[PlanStep(**step) for step in record.steps],
        requires_confirmation=record.requires_confirmation,
        estimated_changes=record.estimated_changes,
        context=record.context,
        created_at=record.created_at,
        expires_at=record.expires_at,
    )


def _tool_results_from_steps(steps: list[dict]) -> list[dict]:
    tool_results: list[dict] = []
    for step in steps:
        tool_results.append(
            {
                "step_number": step.get("step_number"),
                "span_id": step.get("span_id"),
                "parent_span_id": step.get("parent_span_id"),
                "tool_name": step.get("tool_name") or step.get("action") or step.get("step_type"),
                "status": step.get("status") or "completed",
                "latency_ms": step.get("latency_ms", 0),
                "tokens_used": step.get("tokens_used") or 0,
                "observation": step.get("content") or "",
            }
        )
    return tool_results


def _regression_case_payload(execution: AgentExecution) -> dict | None:
    metadata = execution.result_metadata or {}
    case = metadata.get("regression_case")
    if not isinstance(case, dict):
        return None
    payload = dict(case)
    payload.setdefault("regression_case_id", metadata.get("regression_case_id"))
    payload.setdefault("execution_id", str(execution.id))
    payload.setdefault("decision_id", metadata.get("decision_id"))
    payload.setdefault("created_at", metadata.get("feedback_at") or execution.completed_at or execution.created_at)
    payload.setdefault("score", execution.user_feedback)
    return payload


def _export_filename(execution: AgentExecution, suffix: str) -> str:
    stem = f"contract-review-{str(execution.id)[:8]}"
    return f"{stem}.{suffix}"


# 严重度 -> (中文标签, RGB 0-255)
SEVERITY_LABELS: dict[str, tuple[str, tuple[int, int, int]]] = {
    "high": ("高风险", (192, 57, 43)),
    "medium": ("中风险", (202, 138, 4)),
    "low": ("低风险", (30, 126, 52)),
    "uncertain": ("依据不足", (107, 114, 128)),
}
_BASIS_LABELS = {"legal": "法律", "enterprise": "企业制度"}


def _severity_label(value: str | None) -> tuple[str, tuple[int, int, int]]:
    return SEVERITY_LABELS.get(str(value or ""), (str(value or "未知"), (55, 65, 81)))


def _collect_report_refs(report: dict) -> list[dict]:
    seen: set[str] = set()
    refs: list[dict] = []
    for item in report.get("risk_items") or []:
        for ref in item.get("references") or []:
            key = str(ref.get("citation_id") or ref.get("citation_code") or ref.get("chunk_id") or "")
            if not key or key in seen:
                continue
            seen.add(key)
            refs.append(ref)
    return refs


def _report_blocks(execution: AgentExecution) -> list[dict]:
    """把结构化 review_report 拍平成有序排版块；无结构化报告时回退到纯文本。"""
    metadata = execution.result_metadata or {}
    report = metadata.get("review_report") if isinstance(metadata, dict) else None
    blocks: list[dict] = [{"style": "h1", "text": "合同审查报告"}]

    if not isinstance(report, dict) or not report.get("risk_items"):
        for line in (execution.result or "未生成审查结论").splitlines() or ["未生成审查结论"]:
            blocks.append({"style": "body", "text": line})
        return blocks

    overall_label, overall_color = _severity_label(report.get("overall_risk"))
    blocks.append({"style": "label", "text": f"总体风险：{overall_label}", "color": overall_color})
    if report.get("summary"):
        blocks.append({"style": "h2", "text": "审查摘要"})
        blocks.append({"style": "body", "text": str(report["summary"])})

    blocks.append({"style": "h2", "text": "风险项"})
    for idx, item in enumerate(report.get("risk_items") or [], start=1):
        slabel, scolor = _severity_label(item.get("severity"))
        blocks.append({"style": "h3", "text": f"{idx}. [{slabel}] {item.get('issue') or '审查发现'}", "color": scolor})
        for field, key in (
            ("条款摘录", "clause_excerpt"),
            ("法律依据", "legal_basis"),
            ("企业制度依据", "enterprise_basis"),
            ("修改建议", "recommendation"),
        ):
            value = item.get(key)
            if value:
                blocks.append({"style": "field", "label": field, "text": str(value)})
        conf = item.get("confidence")
        if isinstance(conf, (int, float)):
            blocks.append({"style": "field", "label": "置信度", "text": f"{conf * 100:.0f}%"})

    refs = _collect_report_refs(report)
    if refs:
        blocks.append({"style": "h2", "text": "引用依据"})
        for ref in refs:
            kind = _BASIS_LABELS.get(str(ref.get("basis_kind")), "其他")
            code = ref.get("citation_code") or ref.get("doc_title") or ref.get("chunk_id") or ""
            hier = ref.get("hierarchy") or ""
            blocks.append({"style": "ref", "text": f"[{kind}] {code} {hier}".strip()})
    return blocks


def _report_markdown(execution: AgentExecution) -> str:
    metadata = execution.result_metadata or {}
    report = metadata.get("review_report") if isinstance(metadata, dict) else None
    if not isinstance(report, dict) or not report.get("risk_items"):
        return execution.result or ""

    lines = ["# 合同审查报告", "", f"**总体风险：{_severity_label(report.get('overall_risk'))[0]}**", ""]
    if report.get("summary"):
        lines += ["## 审查摘要", "", str(report["summary"]), ""]
    lines += ["## 风险项", ""]
    for idx, item in enumerate(report.get("risk_items") or [], start=1):
        lines.append(f"### {idx}. [{_severity_label(item.get('severity'))[0]}] {item.get('issue') or '审查发现'}")
        lines.append("")
        for field, key in (
            ("条款摘录", "clause_excerpt"),
            ("法律依据", "legal_basis"),
            ("企业制度依据", "enterprise_basis"),
            ("修改建议", "recommendation"),
        ):
            if item.get(key):
                lines.append(f"- **{field}**：{item[key]}")
        conf = item.get("confidence")
        if isinstance(conf, (int, float)):
            lines.append(f"- **置信度**：{conf * 100:.0f}%")
        lines.append("")
    refs = _collect_report_refs(report)
    if refs:
        lines += ["## 引用依据", ""]
        for ref in refs:
            kind = _BASIS_LABELS.get(str(ref.get("basis_kind")), "其他")
            code = ref.get("citation_code") or ref.get("doc_title") or ref.get("chunk_id") or ""
            hier = ref.get("hierarchy") or ""
            lines.append(f"- [{kind}] {code} {hier}".rstrip())
    return "\n".join(lines)


def _docx_bytes(blocks: list[dict]) -> bytes:
    from docx import Document as DocxDocument
    from docx.shared import RGBColor

    doc = DocxDocument()
    for block in blocks:
        style = block.get("style", "body")
        text = block.get("text", "")
        color = block.get("color")
        if style == "h1":
            doc.add_heading(text, level=0)
        elif style == "h2":
            doc.add_heading(text, level=1)
        elif style == "h3":
            heading = doc.add_heading(level=2)
            run = heading.add_run(text)
            if color:
                run.font.color.rgb = RGBColor(*color)
        elif style == "label":
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            run.bold = True
            if color:
                run.font.color.rgb = RGBColor(*color)
        elif style == "field":
            paragraph = doc.add_paragraph()
            label_run = paragraph.add_run(f"{block.get('label', '')}：")
            label_run.bold = True
            paragraph.add_run(text)
        elif style == "ref":
            doc.add_paragraph(text, style="List Bullet")
        else:
            doc.add_paragraph(text)
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


def _pdf_bytes(blocks: list[dict]) -> bytes:
    import fitz

    font = fitz.Font("cjk")  # 内置 Droid Sans Fallback，覆盖中文，修复默认字体中文不显示
    font_name = "DSF"
    page_w, page_h = 595.0, 842.0
    margin_x, margin_top, margin_bottom = 50.0, 56.0, 56.0
    content_w = page_w - 2 * margin_x
    line_height_ratio = 1.4
    # style -> (字号, 段前间距, 默认颜色 0-1)
    style_spec = {
        "h1": (19.0, 0.0, (0.12, 0.16, 0.22)),
        "h2": (14.0, 12.0, (0.12, 0.16, 0.22)),
        "h3": (12.0, 9.0, (0.20, 0.25, 0.32)),
        "label": (12.0, 6.0, (0.12, 0.16, 0.22)),
        "field": (10.5, 3.0, (0.12, 0.16, 0.22)),
        "ref": (10.0, 2.0, (0.30, 0.34, 0.40)),
        "body": (10.5, 3.0, (0.12, 0.16, 0.22)),
    }

    doc = fitz.open()
    state: dict = {"page": None, "y": 0.0}

    def new_page() -> None:
        page = doc.new_page(width=page_w, height=page_h)
        page.insert_font(fontname=font_name, fontbuffer=font.buffer)
        state["page"] = page
        state["y"] = margin_top

    def wrap(text: str, size: float) -> list[str]:
        out: list[str] = []
        for para in (text or "").split("\n"):
            if para == "":
                out.append("")
                continue
            current = ""
            for ch in para:
                if font.text_length(current + ch, fontsize=size) > content_w and current:
                    out.append(current)
                    current = ch
                else:
                    current += ch
            out.append(current)
        return out

    new_page()
    for block in blocks:
        size, gap, default_color = style_spec.get(block.get("style", "body"), style_spec["body"])
        raw = block.get("color")
        color = tuple(c / 255 for c in raw) if raw else default_color
        if block.get("style") == "field":
            display = f"{block.get('label', '')}：{block.get('text', '')}"
        elif block.get("style") == "ref":
            display = f"• {block.get('text', '')}"
        else:
            display = block.get("text", "")

        state["y"] += gap
        line_h = size * line_height_ratio
        for line in wrap(display, size) or [""]:
            if state["y"] + line_h > page_h - margin_bottom:
                new_page()
            state["page"].insert_text(
                fitz.Point(margin_x, state["y"] + size),
                line,
                fontname=font_name,
                fontsize=size,
                color=color,
            )
            state["y"] += line_h

    payload = doc.tobytes()
    doc.close()
    return payload


async def _persist_retrieval_log(
    db: AsyncSession,
    *,
    tenant_id: str,
    session_id: uuid.UUID,
    execution_id: uuid.UUID,
    query: str,
    filters: dict,
    debug: dict,
    context_refs: list[dict],
) -> RetrievalLog:
    row = RetrievalLog(
        id=uuid.uuid4(),
        tenant_id=tenant_id,
        session_id=session_id,
        execution_id=execution_id,
        query=query,
        filters=filters,
        preprocessed=debug.get("preprocessed", {}),
        vector_hits=debug.get("channels", {}).get("vector", []),
        keyword_hits=debug.get("channels", {}).get("keyword", []),
        graph_hits=debug.get("channels", {}).get("graph", []),
        merged_hits=debug.get("merged", []),
        rerank_scores=debug.get("reranked", []),
        filtered_out=debug.get("filtered_out", []),
        final_context=context_refs,
        latency_ms=float(debug.get("latency_ms", 0.0)),
        created_at=datetime.now(timezone.utc),
    )
    db.add(row)
    await db.flush()
    return row


async def _persist_failed_execution(
    db: AsyncSession,
    *,
    req: AgentExecuteRequest,
    ctx: RequestContext,
    execution_id: uuid.UUID,
    trace_id: str,
    start_time: float,
    exc: Exception,
    references: list[dict] | None = None,
    retrieval_debug: dict | None = None,
) -> AgentExecuteResponse:
    completed_at = datetime.now(timezone.utc)
    latency_ms = (time.perf_counter() - start_time) * 1000
    message = f"执行失败：{exc}"
    review_report = (
        _build_contract_review_report(message, references or [], req.query)
        if req.task_type == "contract_review"
        else None
    )
    step_id = uuid.uuid4()
    span_id = uuid.uuid4().hex[:16]
    step_payload = {
        "step_number": 1,
        "step_id": str(step_id),
        "span_id": span_id,
        "parent_span_id": None,
        "step_type": "thought",
        "content": message[:500],
        "action": None,
        "tool_name": None,
        "status": "failed",
        "tokens_used": 0,
        "latency_ms": round(latency_ms, 2),
    }
    debug = retrieval_debug or {}
    execution = AgentExecution(
        id=execution_id,
        trace_id=trace_id,
        session_id=req.session_id,
        user_id=ctx.user_uuid,
        tenant_id=ctx.tenant_id,
        task_type=req.task_type,
        user_query=req.query,
        parsed_intent=(debug.get("preprocessed") or {}).get("intent"),
        parsed_entities=(debug.get("preprocessed") or {}).get("entities", []),
        agent_type=orchestrator_agent.agent_type,
        model_config_id=req.model_config_id,
        prompt_template_id=req.prompt_template_id,
        total_steps=1,
        total_tokens_used=0,
        total_cost=0.0,
        status="failed",
        result=message,
        result_metadata={
            "references": references or [],
            "review_report": review_report,
            "tool_results": _tool_results_from_steps([step_payload]),
            "usage": {
                "total_tokens": 0,
                "retrieval_chunks": len(references or []),
                "retrieval_latency_ms": debug.get("latency_ms", 0.0),
            },
            "agent_metadata": {"error": {"type": type(exc).__name__, "message": str(exc)}},
            "filters": req.filters,
        },
        error_message=message,
        relevance_score=None,
        factuality_score=None,
        user_feedback=None,
        latency_ms=latency_ms,
        retrieval_latency_ms=float(debug.get("latency_ms", 0.0)),
        generation_latency_ms=0.0,
        created_at=completed_at,
        completed_at=completed_at,
    )
    db.add(execution)
    db.add(
        AgentStep(
            id=step_id,
            execution_id=execution_id,
            trace_id=trace_id,
            span_id=span_id,
            parent_span_id=None,
            step_number=1,
            step_type="thought",
            agent_type=orchestrator_agent.agent_type,
            thought=message,
            action=None,
            action_input={},
            observation=message,
            tool_name=None,
            tool_input={},
            tool_output=message,
            retrieved_chunks=[],
            retrieval_scores=[],
            tokens_used=0,
            latency_ms=latency_ms,
            status="failed",
            error_message=message,
            started_at=completed_at,
            completed_at=completed_at,
        )
    )
    await session_memory_service.append_message(
        db=db,
        session_id=req.session_id,
        tenant_id=ctx.tenant_id,
        role="assistant",
        content=message,
        user_id=ctx.user_uuid,
        trace_id=trace_id,
        metadata={"execution_id": str(execution_id), "task_type": req.task_type, "status": "failed"},
    )
    await db.flush()
    return AgentExecuteResponse(
        execution_id=execution.id,
        trace_id=execution.trace_id,
        status=execution.status,
        result=execution.result or "",
        references=references or [],
        steps=[step_payload],
        usage=execution.result_metadata["usage"],
        latency_ms=round(latency_ms, 2),
        review_report=review_report,
        tool_results=execution.result_metadata["tool_results"],
    )


@router.post("/plan", response_model=AgentPlanResponse)
async def plan_agent(
    req: AgentPlanRequest,
    db: AsyncSession = Depends(get_write_db),
    ctx: RequestContext = Depends(get_request_context),
):
    """Create and persist a confirmable plan without creating an execution record."""

    if req.tenant_id != ctx.tenant_id:
        raise HTTPException(status_code=403, detail="tenant mismatch")

    now = datetime.now(timezone.utc)
    await _expire_stale_decisions(db, now)
    decision_id = f"dec_{uuid.uuid4().hex[:24]}"
    plan = _build_agent_plan(req)
    record = AgentDecision(
        decision_id=decision_id,
        tenant_id=ctx.tenant_id,
        user_id=ctx.user_id,
        session_id=req.session_id,
        query=req.query,
        task_type=req.task_type,
        filters=plan["context"].get("filters", {}),
        intent_summary=plan["intent_summary"],
        steps=plan["steps"],
        requires_confirmation=plan["requires_confirmation"],
        estimated_changes=plan["estimated_changes"],
        context=plan["context"],
        status="planned",
        user_confirmation={},
        created_at=now,
        expires_at=now + timedelta(seconds=DECISION_TTL_SECONDS),
        updated_at=now,
    )
    db.add(record)
    await db.flush()
    return _decision_to_plan_response(record)


@router.get("/tool-catalog")
async def get_tool_catalog():
    return {
        "version": "2026-05-14",
        "domains": tool_catalog_by_domain(),
    }


@router.get("/workspace-extension-schema")
async def get_workspace_extension_schema():
    return {
        "version": "2026-05-14",
        "status": "reserved",
        "schemas": {
            "contract_matter": ContractMatter.model_json_schema(),
            "performance_event": PerformanceEvent.model_json_schema(),
            "expense_evidence": ExpenseEvidence.model_json_schema(),
        },
    }


@router.post("/decisions/{decision_id}/execute", response_model=AgentExecuteResponse)
async def execute_decision(
    decision_id: str,
    confirmation: AgentDecisionExecuteRequest,
    db: AsyncSession = Depends(get_write_db),
    ctx: RequestContext = Depends(get_request_context),
):
    """Execute a previously generated plan after user confirmation."""

    await _expire_stale_decisions(db)
    record = await db.scalar(
        select(AgentDecision).where(
            AgentDecision.decision_id == decision_id,
        )
    )
    if not record:
        raise HTTPException(status_code=404, detail="decision not found or expired")
    if record.tenant_id != ctx.tenant_id:
        raise HTTPException(status_code=403, detail="tenant mismatch")
    if record.status == "expired":
        raise HTTPException(status_code=404, detail="decision not found or expired")
    if record.status == "executed":
        raise HTTPException(status_code=409, detail="decision already executed")
    if record.requires_confirmation and not confirmation.confirmed:
        raise HTTPException(status_code=409, detail="decision requires user confirmation")

    execute_req = AgentExecuteRequest(
        query=record.query,
        task_type=record.task_type or "contract_review",
        session_id=record.session_id,
        tenant_id=ctx.tenant_id,
        filters=record.filters or {},
    )
    response = await execute_agent(execute_req, db=db, ctx=ctx)
    tool_results = _tool_results_from_steps(response.steps)
    plan_payload = {
        "decision_id": decision_id,
        "intent_summary": record.intent_summary,
        "steps": record.steps,
        "requires_confirmation": record.requires_confirmation,
        "estimated_changes": record.estimated_changes,
        "context": record.context,
        "created_at": record.created_at.isoformat(),
        "expires_at": record.expires_at.isoformat(),
    }
    user_confirmation = {
        "confirmed": confirmation.confirmed,
        "comment": confirmation.comment,
        "confirmed_at": datetime.now(timezone.utc).isoformat(),
        "user_id": ctx.user_id,
    }

    execution = await db.scalar(
        select(AgentExecution).where(
            AgentExecution.id == response.execution_id,
            AgentExecution.tenant_id == ctx.tenant_id,
        )
    )
    if execution:
        metadata = dict(execution.result_metadata or {})
        metadata.update(
            {
                "decision_id": decision_id,
                "plan": plan_payload,
                "tool_results": tool_results,
                "user_confirmation": user_confirmation,
            }
        )
        execution.result_metadata = metadata
        await db.flush()

    now = datetime.now(timezone.utc)
    record.status = "executed"
    record.execution_id = response.execution_id
    record.trace_id = response.trace_id
    record.user_confirmation = user_confirmation
    record.updated_at = now
    response.decision_id = decision_id
    response.plan = plan_payload
    response.tool_results = tool_results
    return response


@router.post("/execute", response_model=AgentExecuteResponse)
async def execute_agent(
    req: AgentExecuteRequest,
    db: AsyncSession = Depends(get_write_db),
    ctx: RequestContext = Depends(get_request_context),
):
    if req.tenant_id != ctx.tenant_id:
        raise HTTPException(status_code=403, detail="tenant mismatch")

    await session_memory_service.ensure_session(
        db=db,
        session_id=req.session_id,
        tenant_id=ctx.tenant_id,
        user_id=ctx.user_uuid,
    )

    user_message = await session_memory_service.append_message(
        db=db,
        session_id=req.session_id,
        tenant_id=ctx.tenant_id,
        role="user",
        content=req.query,
        user_id=ctx.user_uuid,
    )
    await session_memory_service.upsert_facts_from_message(
        db=db,
        session_id=req.session_id,
        tenant_id=ctx.tenant_id,
        source_message_id=user_message.id,
        text=req.query,
    )

    execution_id = uuid.uuid4()
    trace_id = uuid.uuid4().hex[:32]
    start_time = time.perf_counter()

    try:
        retrieval_results, retrieval_debug = await hybrid_retriever.retrieve_with_debug(
            query=req.query,
            tenant_id=ctx.tenant_id,
            filters=req.filters,
        )
    except Exception as exc:
        return await _persist_failed_execution(
            db,
            req=req,
            ctx=ctx,
            execution_id=execution_id,
            trace_id=trace_id,
            start_time=start_time,
            exc=exc,
        )

    if req.task_type == "contract_review":
        retrieval_results = await _merge_enterprise_rules(
            query=req.query,
            tenant_id=ctx.tenant_id,
            base_results=retrieval_results,
        )

    built_context = context_builder.build(retrieval_results)
    # Create citation records and attach to references.
    citation_by_chunk: dict[str, dict] = {}
    for item in retrieval_results:
        citation = await citation_service.ensure_for_result(
            db=db,
            tenant_id=ctx.tenant_id,
            session_id=req.session_id,
            execution_id=execution_id,
            chunk_id=item.chunk_id,
            document_id=item.metadata.get("doc_id"),
            excerpt=item.content[:1200],
            title=item.metadata.get("doc_title"),
            locator=item.metadata.get("hierarchy_path"),
            metadata=item.metadata,
        )
        citation_by_chunk[item.chunk_id] = {
            "citation_id": str(citation.id),
            "citation_code": citation.citation_code,
        }

    for ref in built_context["references"]:
        citation_info = citation_by_chunk.get(ref.get("chunk_id"))
        if citation_info:
            ref.update(citation_info)

    memory_ctx = await session_memory_service.get_runtime_context(
        db=db,
        session_id=req.session_id,
        tenant_id=ctx.tenant_id,
    )
    profile_ctx = await session_memory_service.get_profile_context(
        db=db,
        user_id=ctx.user_uuid,
        tenant_id=ctx.tenant_id,
    )
    await session_memory_service.extract_profile_from_text(
        db=db,
        user_id=ctx.user_uuid,
        tenant_id=ctx.tenant_id,
        text=req.query,
    )
    agent_context = {
        "tenant_id": ctx.tenant_id,
        "retrieval_context": built_context["context"],
        "references": built_context["references"],
        "conversation_history": memory_ctx["history_text"],
        "session_summary": memory_ctx["summary"],
        "memory_facts": memory_ctx["facts"],
        "user_profile": profile_ctx.get("profile_text", ""),
    }

    generation_start = time.perf_counter()
    try:
        result = await orchestrator_agent.execute(req.query, context=agent_context)
    except Exception as exc:
        return await _persist_failed_execution(
            db,
            req=req,
            ctx=ctx,
            execution_id=execution_id,
            trace_id=trace_id,
            start_time=start_time,
            exc=exc,
            references=built_context["references"],
            retrieval_debug=retrieval_debug,
        )
    generation_latency_ms = (time.perf_counter() - generation_start) * 1000

    relevance_score = None
    factuality_score = None
    evaluation_error = None
    if result.success and built_context["references"]:
        try:
            validation_messages = [
                {
                    "role": "system",
                    "content": (
                        "Score the following agent output on two dimensions:\n"
                        "1) Relevance to the user query (0.0-1.0)\n"
                        "2) Factuality — groundedness in the provided references (0.0-1.0)\n"
                        "Output JSON only: {\"relevance\": <float>, \"factuality\": <float>}"
                    ),
                },
                {
                    "role": "user",
                    "content": f"Query: {req.query[:500]}\n\nOutput: {result.output[:2000]}\n\nReferences: {str(built_context['references'][:5])[:2000]}",
                },
            ]
            score_result = await llm_service.light_generate(validation_messages)
            import re as _re
            score_text = score_result.get("content", "{}")
            _match = _re.search(r'\{[^}]+\}', score_text)
            if _match:
                import json as _json
                scores = _json.loads(_match.group())
                relevance_score = scores.get("relevance")
                factuality_score = scores.get("factuality")
        except Exception as exc:
            evaluation_error = {"type": type(exc).__name__, "message": str(exc)}
            logger.debug(f"agent output evaluation failed trace_id={trace_id}: {exc}")
    latency_ms = (time.perf_counter() - start_time) * 1000
    completed_at = datetime.now(timezone.utc)
    review_report = None
    if req.task_type == "contract_review":
        review_report = _build_contract_review_report(result.output or "", built_context["references"], req.query)

    execution = AgentExecution(
        id=execution_id,
        trace_id=trace_id,
        session_id=req.session_id,
        user_id=ctx.user_uuid,
        tenant_id=ctx.tenant_id,
        task_type=req.task_type,
        user_query=req.query,
        parsed_intent=(retrieval_debug.get("preprocessed") or {}).get("intent"),
        parsed_entities=(retrieval_debug.get("preprocessed") or {}).get("entities", []),
        agent_type=orchestrator_agent.agent_type,
        model_config_id=req.model_config_id,
        prompt_template_id=req.prompt_template_id,
        total_steps=len(result.steps),
        total_tokens_used=result.total_tokens,
        total_cost=0.0,
        status="completed" if result.success else "failed",
        result=result.output,
        result_metadata={
            "references": built_context["references"],
            "review_report": review_report,
            "usage": {
                "total_tokens": result.total_tokens,
                "retrieval_chunks": built_context["chunk_count"],
                "retrieval_latency_ms": retrieval_debug.get("latency_ms", 0.0),
            },
            "agent_metadata": {**(result.metadata or {}), "evaluation_error": evaluation_error},
            "filters": req.filters,
        },
        error_message=None if result.success else result.output,
        relevance_score=relevance_score,
        factuality_score=factuality_score,
        user_feedback=None,
        latency_ms=latency_ms,
        retrieval_latency_ms=float(retrieval_debug.get("latency_ms", 0.0)),
        generation_latency_ms=generation_latency_ms,
        created_at=completed_at,
        completed_at=completed_at,
    )
    db.add(execution)

    await _persist_retrieval_log(
        db=db,
        tenant_id=ctx.tenant_id,
        session_id=req.session_id,
        execution_id=execution_id,
        query=req.query,
        filters=req.filters,
        debug=retrieval_debug,
        context_refs=built_context["references"],
    )

    step_rows: list[AgentStep] = []
    step_payloads: list[dict] = []
    for idx, step in enumerate(result.steps):
        step_started = datetime.fromtimestamp(step.timestamp, tz=timezone.utc)
        step_completed = step_started + timedelta(milliseconds=float(step.latency_ms or 0.0))
        step_status = "completed" if result.success or idx < len(result.steps) - 1 else "failed"
        try:
            step_id = uuid.UUID(step.id)
        except ValueError:
            step_id = uuid.uuid4()
        span_id = uuid.uuid4().hex[:16]

        step_rows.append(
            AgentStep(
                id=step_id,
                execution_id=execution_id,
                trace_id=trace_id,
                span_id=span_id,
                parent_span_id=None,
                step_number=idx + 1,
                step_type=step.step_type.value,
                agent_type=orchestrator_agent.agent_type,
                thought=step.content if step.step_type.value == "thought" else None,
                action=step.action,
                action_input=step.action_input or {},
                observation=step.observation or (step.content if step.step_type.value != "action" else None),
                tool_name=step.tool_name,
                tool_input=step.action_input or {},
                tool_output=step.observation,
                retrieved_chunks=[],
                retrieval_scores=[],
                tokens_used=step.tokens_used or 0,
                latency_ms=float(step.latency_ms or 0.0),
                status=step_status,
                error_message=None if step_status == "completed" else result.output,
                started_at=step_started,
                completed_at=step_completed,
            )
        )
        step_payloads.append(
            {
                "step_number": idx + 1,
                "step_id": str(step_id),
                "span_id": span_id,
                "parent_span_id": None,
                "step_type": step.step_type.value,
                "content": (step.content or "")[:500],
                "action": step.action,
                "tool_name": step.tool_name,
                "status": step_status,
                "tokens_used": step.tokens_used,
                "latency_ms": round(float(step.latency_ms or 0.0), 2),
            }
        )

    execution_metadata = dict(execution.result_metadata or {})
    execution_metadata["tool_results"] = _tool_results_from_steps(step_payloads)
    execution.result_metadata = execution_metadata
    db.add_all(step_rows)

    await session_memory_service.append_message(
        db=db,
        session_id=req.session_id,
        tenant_id=ctx.tenant_id,
        role="assistant",
        content=result.output or "",
        user_id=ctx.user_uuid,
        trace_id=trace_id,
        metadata={"execution_id": str(execution_id), "task_type": req.task_type},
    )
    await session_memory_service.refresh_rolling_summary(
        db=db,
        session_id=req.session_id,
        tenant_id=ctx.tenant_id,
    )

    await db.flush()
    return AgentExecuteResponse(
        execution_id=execution.id,
        trace_id=execution.trace_id,
        status=execution.status,
        result=execution.result or "",
        references=built_context["references"],
        steps=step_payloads,
        usage={
            "total_tokens": result.total_tokens,
            "retrieval_chunks": built_context["chunk_count"],
            "retrieval_latency_ms": retrieval_debug.get("latency_ms", 0.0),
        },
        latency_ms=round(latency_ms, 2),
        review_report=review_report,
        tool_results=execution_metadata["tool_results"],
    )


@router.post("/chat")
async def chat_stream(
    req: AgentExecuteRequest,
    db: AsyncSession = Depends(get_write_db),
    ctx: RequestContext = Depends(get_request_context),
):
    if req.tenant_id != ctx.tenant_id:
        raise HTTPException(status_code=403, detail="tenant mismatch")

    await session_memory_service.ensure_session(
        db=db,
        session_id=req.session_id,
        tenant_id=ctx.tenant_id,
        user_id=ctx.user_uuid,
    )
    user_message = await session_memory_service.append_message(
        db=db,
        session_id=req.session_id,
        tenant_id=ctx.tenant_id,
        role="user",
        content=req.query,
        user_id=ctx.user_uuid,
    )
    await session_memory_service.upsert_facts_from_message(
        db=db,
        session_id=req.session_id,
        tenant_id=ctx.tenant_id,
        source_message_id=user_message.id,
        text=req.query,
    )

    trace_id = uuid.uuid4().hex[:32]
    memory_ctx = await session_memory_service.get_runtime_context(
        db=db,
        session_id=req.session_id,
        tenant_id=ctx.tenant_id,
    )

    retrieval_results, _retrieval_debug = await hybrid_retriever.retrieve_with_debug(
        query=req.query,
        tenant_id=ctx.tenant_id,
        filters=req.filters,
    )
    built_context = context_builder.build(retrieval_results)

    messages = [{"role": "system", "content": orchestrator_agent._build_system_prompt()}]
    if built_context["context"]:
        messages.append({"role": "system", "content": f"Retrieval context:\n{built_context['context']}"})
    if memory_ctx["summary"]:
        messages.append({"role": "system", "content": f"Session summary:\n{memory_ctx['summary']}"})
    if memory_ctx["history_text"]:
        messages.append({"role": "system", "content": f"Recent conversation:\n{memory_ctx['history_text']}"})
    messages.append({"role": "user", "content": req.query})

    async def generate():
        chunks: list[str] = []
        async for chunk in llm_service.generate_stream(messages):
            chunks.append(chunk)
            yield f"data: {json.dumps({'type': 'content', 'text': chunk}, ensure_ascii=False)}\n\n"

        final_text = "".join(chunks).strip()
        if final_text:
            await session_memory_service.append_message(
                db=db,
                session_id=req.session_id,
                tenant_id=ctx.tenant_id,
                role="assistant",
                content=final_text,
                user_id=ctx.user_uuid,
                trace_id=trace_id,
                metadata={"stream": True, "references": built_context["references"][:10]},
            )
            await session_memory_service.refresh_rolling_summary(
                db=db,
                session_id=req.session_id,
                tenant_id=ctx.tenant_id,
            )

        refs_data = json.dumps({"type": "references", "data": built_context["references"][:10]}, ensure_ascii=False)
        yield f"data: {refs_data}\n\n"
        yield "data: [DONE]\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")


@router.get("/executions")
async def list_executions(
    task_type: str = Query(default=""),
    status: str = Query(default=""),
    session_id: str = Query(default=""),
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=20, ge=1, le=100),
    db: AsyncSession = Depends(get_read_db),
    ctx: RequestContext = Depends(get_request_context),
):
    query = select(AgentExecution).where(AgentExecution.tenant_id == ctx.tenant_id)
    if task_type:
        query = query.where(AgentExecution.task_type == task_type)
    if status:
        query = query.where(AgentExecution.status == status)
    if session_id:
        query = query.where(AgentExecution.session_id == _parse_uuid(session_id, "session_id"))

    total_stmt = select(func.count()).select_from(query.subquery())
    total = int((await db.scalar(total_stmt)) or 0)
    rows = (
        await db.scalars(
            query.order_by(AgentExecution.created_at.desc())
            .offset((page - 1) * page_size)
            .limit(page_size)
        )
    ).all()

    return {
        "items": [_execution_payload(row) for row in rows],
        "total": total,
        "page": page,
        "page_size": page_size,
        "total_pages": (total + page_size - 1) // page_size,
    }


@router.get("/regression-cases")
async def list_regression_cases(
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=20, ge=1, le=100),
    db: AsyncSession = Depends(get_read_db),
    ctx: RequestContext = Depends(get_request_context),
):
    rows = (
        await db.scalars(
            select(AgentExecution)
            .where(AgentExecution.tenant_id == ctx.tenant_id)
            .order_by(AgentExecution.created_at.desc())
            .limit(2000)
        )
    ).all()
    cases = [payload for row in rows if (payload := _regression_case_payload(row))]
    start = (page - 1) * page_size
    items = cases[start : start + page_size]
    return {
        "items": items,
        "total": len(cases),
        "page": page,
        "page_size": page_size,
        "total_pages": (len(cases) + page_size - 1) // page_size,
    }


@router.get("/regression-cases/{regression_case_id}")
async def get_regression_case(
    regression_case_id: str,
    db: AsyncSession = Depends(get_read_db),
    ctx: RequestContext = Depends(get_request_context),
):
    rows = (
        await db.scalars(
            select(AgentExecution)
            .where(AgentExecution.tenant_id == ctx.tenant_id)
            .order_by(AgentExecution.created_at.desc())
            .limit(2000)
        )
    ).all()
    for row in rows:
        payload = _regression_case_payload(row)
        if payload and payload.get("regression_case_id") == regression_case_id:
            return payload
    raise HTTPException(status_code=404, detail="regression case not found")


@router.get("/executions/{execution_id}")
async def get_execution_detail(
    execution_id: str,
    db: AsyncSession = Depends(get_read_db),
    ctx: RequestContext = Depends(get_request_context),
):
    execution_uuid = _parse_uuid(execution_id, "execution_id")
    execution = await db.scalar(
        select(AgentExecution).where(
            AgentExecution.id == execution_uuid,
            AgentExecution.tenant_id == ctx.tenant_id,
        )
    )
    if not execution:
        raise HTTPException(status_code=404, detail="execution not found")

    step_rows = (
        await db.scalars(
            select(AgentStep)
            .where(AgentStep.execution_id == execution_uuid)
            .order_by(AgentStep.step_number.asc())
        )
    ).all()
    return _execution_payload(execution, steps=[_step_payload(step) for step in step_rows])


@router.get("/executions/{execution_id}/export")
async def export_execution_report(
    execution_id: str,
    format: str = Query(default="markdown", pattern="^(markdown|docx|pdf)$"),
    db: AsyncSession = Depends(get_read_db),
    ctx: RequestContext = Depends(get_request_context),
):
    execution_uuid = _parse_uuid(execution_id, "execution_id")
    execution = await db.scalar(
        select(AgentExecution).where(
            AgentExecution.id == execution_uuid,
            AgentExecution.tenant_id == ctx.tenant_id,
        )
    )
    if not execution:
        raise HTTPException(status_code=404, detail="execution not found")

    if format == "docx":
        payload = _docx_bytes(_report_blocks(execution))
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = _export_filename(execution, "docx")
    elif format == "pdf":
        payload = _pdf_bytes(_report_blocks(execution))
        media_type = "application/pdf"
        filename = _export_filename(execution, "pdf")
    else:
        payload = _report_markdown(execution).encode("utf-8")
        media_type = "text/markdown; charset=utf-8"
        filename = _export_filename(execution, "md")

    return StreamingResponse(
        io.BytesIO(payload),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/trace/{trace_id}")
async def get_trace(
    trace_id: str,
    db: AsyncSession = Depends(get_read_db),
    ctx: RequestContext = Depends(get_request_context),
):
    execution = await db.scalar(
        select(AgentExecution).where(
            AgentExecution.trace_id == trace_id,
            AgentExecution.tenant_id == ctx.tenant_id,
        )
    )
    if not execution:
        raise HTTPException(status_code=404, detail="trace not found")

    step_rows = (
        await db.scalars(
            select(AgentStep)
            .where(AgentStep.execution_id == execution.id)
            .order_by(AgentStep.step_number.asc())
        )
    ).all()
    return _execution_payload(execution, steps=[_step_payload(step) for step in step_rows])


@router.post("/executions/{execution_id}/feedback")
async def submit_feedback(
    execution_id: str,
    score: int,
    comment: str = "",
    db: AsyncSession = Depends(get_write_db),
    ctx: RequestContext = Depends(get_request_context),
):
    if not 1 <= score <= 5:
        raise HTTPException(status_code=400, detail="score must be between 1 and 5")

    execution_uuid = _parse_uuid(execution_id, "execution_id")
    execution = await db.scalar(
        select(AgentExecution).where(
            AgentExecution.id == execution_uuid,
            AgentExecution.tenant_id == ctx.tenant_id,
        )
    )
    if not execution:
        raise HTTPException(status_code=404, detail="execution not found")

    metadata = dict(execution.result_metadata or {})
    metadata["user_comment"] = comment
    metadata["feedback_at"] = datetime.now(timezone.utc).isoformat()
    metadata["feedback_user_id"] = ctx.user_id
    metadata["user_feedback"] = score
    regression_case_id = metadata.get("regression_case_id") or f"reg_{uuid.uuid4().hex[:16]}"
    metadata["regression_case_id"] = regression_case_id
    metadata["regression_case"] = {
        "regression_case_id": regression_case_id,
        "created_at": metadata["feedback_at"],
        "tenant_id": ctx.tenant_id,
        "user_id": ctx.user_id,
        "execution_id": str(execution.id),
        "decision_id": metadata.get("decision_id"),
        "input": {
            "query": execution.user_query,
            "task_type": execution.task_type,
            "session_id": str(execution.session_id),
        },
        "plan": metadata.get("plan"),
        "tool_results": metadata.get("tool_results", []),
        "review_report": metadata.get("review_report"),
        "references": metadata.get("references", []),
        "expected_correction": comment,
        "score": score,
    }
    execution.user_feedback = score
    execution.result_metadata = metadata
    await db.flush()
    return {"message": "feedback submitted", "regression_case_id": regression_case_id}

